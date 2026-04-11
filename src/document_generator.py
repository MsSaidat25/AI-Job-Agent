"""
Copyright 2026 AVIEN SOLUTIONS INC (www.aviensolutions.com).
All Rights Reserved.
No part of this software or any of its contents may be reproduced, copied,
modified or adapted, without the prior written consent of the author, unless
otherwise indicated for stand-alone materials.
For permission requests, write to the publisher at the email address below:
avien@aviensolutions.com
THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN
THE SOFTWARE.

Document Generator — Resume & Cover Letter.

Design
──────
• Claude is used to produce a first draft tailored to each job listing.
• The output is plain Markdown (easy to convert to PDF/DOCX downstream).
• PII is passed to the generator because the *user* controls what goes in
  their own documents; it is never sent to third-party job boards.
• A tailoring_notes field explains every significant edit so the user
  understands why choices were made.

Customisation hooks
───────────────────
• RESUME_TONE  — "professional" (default) | "creative" | "technical"
• Extend DocumentGenerator with a custom template by subclassing and
  overriding _resume_system_prompt() / _cover_letter_system_prompt().
"""
from __future__ import annotations

import json
import re
import uuid
from datetime import datetime, timezone
from io import BytesIO
from typing import Any, cast

from anthropic import Anthropic
from anthropic.types import TextBlock

from config.settings import AGENT_MODEL, MAX_TOKENS
from src.llm_client import create_message_with_failover, get_llm_client
from src.models import GeneratedDocument, JobListing, UserProfile
from src.utils import strip_json_fences


_RESUME_SYSTEM = """You are an expert resume writer with 15+ years of experience
across tech, finance, marketing, and creative industries worldwide.

Rules:
1. Write in Markdown. Use ## for section headers, bold for company/role names.
2. Tailor bullet points to mirror the job listing's keywords and requirements.
3. Quantify achievements wherever possible (e.g. "reduced latency by 40%").
4. Keep to 1–2 pages of content (roughly 500–800 words of body text).
5. Never invent credentials or experiences not supplied by the user.
6. Do NOT include any protected-attribute language (age, gender, etc.).
7. End with a JSON block between ```json ``` tags containing "tailoring_notes".
"""

_COVER_LETTER_SYSTEM = """You are a professional career coach who writes compelling,
personalised cover letters that stand out without being gimmicky.

Rules:
1. Write in plain prose (no bullet lists in the body).
2. Opening: hook with a specific reason why you're excited about THIS company.
3. Middle: bridge 2–3 key skills/experiences to the job's core requirements.
4. Closing: confident call to action, no desperate begging.
5. Tone: enthusiastic but professional. Adapt to regional/industry norms.
6. Length: 3–4 paragraphs (~300–400 words).
7. End with a JSON block between ```json ``` tags containing "tailoring_notes".
"""


class DocumentGenerator:
    def __init__(self, client: Anthropic | None = None) -> None:
        self._client = client or get_llm_client()

    # ── Public API ─────────────────────────────────────────────────────────

    def generate_cover_letter(
        self,
        profile: UserProfile,
        job: JobListing,
    ) -> GeneratedDocument:
        """Generate a tailored cover letter for *job*."""
        user_prompt = self._build_cover_letter_prompt(profile, job)
        raw = self._call_model(_COVER_LETTER_SYSTEM, user_prompt)
        content, notes = self._split_notes(raw)
        return GeneratedDocument(
            id=str(uuid.uuid4()),
            user_id=profile.id,
            job_id=job.id,
            doc_type="cover_letter",
            content=content,
            created_at=datetime.now(timezone.utc),
            model_used=AGENT_MODEL,
            tailoring_notes=notes,
        )

    def score_ats_match(
        self,
        resume_text: str,
        job_description: str,
    ) -> dict[str, Any]:
        """Score how well a resume matches a job description for ATS systems.

        Returns dict with ats_score (0-100), missing_keywords, and suggestions.
        """
        prompt = f"""Analyse the resume against the job description for ATS (Applicant Tracking System) compatibility.

--- RESUME ---
{resume_text[:3000]}

--- JOB DESCRIPTION ---
{job_description[:3000]}

Return ONLY valid JSON (no markdown fences) with these keys:
- "ats_score": integer 0-100 representing match percentage
- "missing_keywords": list of important keywords from the JD missing in the resume
- "matched_keywords": list of keywords found in both
- "suggestions": list of 3-5 specific insertions to improve the score
"""
        try:
            response = create_message_with_failover(
                self._client,
                model=AGENT_MODEL,
                max_tokens=1024,
                system="You are an ATS optimisation expert. Respond ONLY with valid JSON.",
                messages=[{"role": "user", "content": prompt}],
            )
            text = cast(TextBlock, response.content[0]).text.strip()
            data = json.loads(strip_json_fences(text))
            score = max(0, min(100, int(data.get("ats_score", 0))))
            return {
                "ats_score": score,
                "missing_keywords": data.get("missing_keywords", []),
                "matched_keywords": data.get("matched_keywords", []),
                "suggestions": data.get("suggestions", []),
            }
        except Exception:
            return {
                "ats_score": 0,
                "missing_keywords": [],
                "matched_keywords": [],
                "suggestions": ["Unable to compute ATS score. Please try again."],
            }

    def generate_resume(
        self,
        profile: UserProfile,
        job: JobListing,
        tone: str = "professional",
        auto_ats: bool = True,
    ) -> GeneratedDocument:
        """Generate a tailored resume for *job* based on *profile*.

        When auto_ats is True, automatically scores the resume against the JD.
        """
        user_prompt = self._build_resume_prompt(profile, job, tone)
        raw = self._call_model(_RESUME_SYSTEM, user_prompt)
        content, notes = self._split_notes(raw)

        ats_score: float | None = None
        missing_keywords: list[str] = []
        if auto_ats and job.description:
            ats_result = self.score_ats_match(content, job.description)
            ats_score = ats_result["ats_score"]
            missing_keywords = ats_result["missing_keywords"]

        return GeneratedDocument(
            id=str(uuid.uuid4()),
            user_id=profile.id,
            job_id=job.id,
            doc_type="resume",
            content=content,
            created_at=datetime.now(timezone.utc),
            model_used=AGENT_MODEL,
            tailoring_notes=notes,
            ats_score=ats_score,
            missing_keywords=missing_keywords,
        )

    def suggest_improvements(
        self,
        document: GeneratedDocument,
        job: JobListing,
    ) -> str:
        """Return a bullet-list of specific improvement suggestions."""
        prompt = f"""Review the following {document.doc_type} for a {job.title} role
at {job.company}.

--- DOCUMENT ---
{document.content}

--- JOB REQUIREMENTS ---
{json.dumps(job.requirements)}

Provide 5–7 specific, actionable improvement suggestions as a numbered list.
Focus on: keyword optimisation, impact quantification, relevance, and structure.
"""
        response = create_message_with_failover(
            self._client,
            model=AGENT_MODEL,
            max_tokens=700,
            messages=[{"role": "user", "content": prompt}],
        )
        return cast(TextBlock, response.content[0]).text.strip()

    def export_document(self, content_md: str, format: str) -> bytes:  # noqa: A002
        """Export markdown content as PDF or DOCX bytes.

        Args:
            content_md: Markdown-formatted document content.
            format: ``"pdf"`` or ``"docx"``.

        Returns:
            Raw file bytes ready for download.

        Raises:
            ValueError: If *format* is not ``"pdf"`` or ``"docx"``.
        """
        if format == "pdf":
            return self._export_pdf(content_md)
        if format == "docx":
            return self._export_docx(content_md)
        raise ValueError(f"Unsupported format: {format!r}. Must be 'pdf' or 'docx'.")

    # ── Private helpers ────────────────────────────────────────────────────

    def _sanitise_for_pdf(self, text: str) -> str:
        """Replace characters outside latin-1 so fpdf2 core fonts don't crash."""
        return text.encode("latin-1", errors="replace").decode("latin-1")

    def _export_pdf(self, content_md: str) -> bytes:
        from fpdf import FPDF  # lazy import — optional dependency

        pdf = FPDF()
        pdf.set_margins(20, 20, 20)
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)

        for line in content_md.splitlines():
            stripped = line.strip()
            if stripped.startswith("# ") and not stripped.startswith("## "):
                pdf.set_font("Helvetica", "B", 16)
                pdf.ln(4)
                pdf.cell(0, 10, self._sanitise_for_pdf(stripped[2:]), new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)
            elif stripped.startswith("## "):
                pdf.set_font("Helvetica", "B", 13)
                pdf.ln(3)
                pdf.cell(0, 8, self._sanitise_for_pdf(stripped[3:]), new_x="LMARGIN", new_y="NEXT")
                pdf.ln(1)
            elif stripped.startswith("### "):
                pdf.set_font("Helvetica", "B", 11)
                pdf.ln(2)
                pdf.cell(0, 7, self._sanitise_for_pdf(stripped[4:]), new_x="LMARGIN", new_y="NEXT")
            elif stripped.startswith(("- ", "* ")):
                pdf.set_font("Helvetica", "", 10)
                text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped[2:])
                pdf.multi_cell(0, 6, f"  - {self._sanitise_for_pdf(text)}", new_x="LMARGIN", new_y="NEXT")
            elif stripped == "":
                pdf.ln(3)
            else:
                pdf.set_font("Helvetica", "", 10)
                text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
                pdf.multi_cell(0, 6, self._sanitise_for_pdf(text), new_x="LMARGIN", new_y="NEXT")

        return bytes(pdf.output())

    def _export_docx(self, content_md: str) -> bytes:
        from docx import Document  # lazy import — optional dependency
        from docx.shared import Pt

        doc = Document()

        # Tighten default paragraph spacing
        style = doc.styles["Normal"]
        style.font.size = Pt(11)  # type: ignore[union-attr]

        for line in content_md.splitlines():
            stripped = line.strip()
            if stripped.startswith("# ") and not stripped.startswith("## "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith(("- ", "* ")):
                text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped[2:])
                doc.add_paragraph(text, style="List Bullet")
            elif stripped == "":
                continue
            else:
                # Render inline **bold** spans
                para = doc.add_paragraph()
                for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", stripped)):
                    run = para.add_run(part)
                    run.bold = (i % 2 == 1)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _call_model(self, system: str, user_content: str) -> str:
        response = create_message_with_failover(
            self._client,
            model=AGENT_MODEL,
            max_tokens=MAX_TOKENS,
            system=system,
            messages=[{"role": "user", "content": user_content}],
        )
        return cast(TextBlock, response.content[0]).text.strip()

    def _split_notes(self, raw: str) -> tuple[str, str]:
        """Separate document body from trailing ```json ... ``` tailoring notes."""
        if "```json" in raw:
            body, _, rest = raw.partition("```json")
            notes_raw = rest.partition("```")[0].strip()
            try:
                notes_data = json.loads(notes_raw)
                notes = notes_data.get("tailoring_notes", notes_raw)
            except json.JSONDecodeError:
                notes = notes_raw
            return body.strip(), notes
        return raw, ""

    def _build_resume_prompt(
        self, profile: UserProfile, job: JobListing, tone: str
    ) -> str:
        edu_text = json.dumps(profile.education, indent=2) if profile.education else "Not provided"
        work_text = json.dumps(profile.work_history, indent=2) if profile.work_history else "Not provided"
        return f"""Create a {tone} resume for the following candidate applying to the job below.

=== CANDIDATE PROFILE ===
Name: {profile.name}
Location: {profile.location}
Email: {profile.email}
Skills: {", ".join(profile.skills)}
Experience Level: {profile.experience_level.value} ({profile.years_of_experience} years)
Languages: {", ".join(profile.languages)}
Certifications: {", ".join(profile.certifications) or "None"}
Portfolio: {profile.portfolio_url or "N/A"}
LinkedIn: {profile.linkedin_url or "N/A"}

Education:
{edu_text}

Work History:
{work_text}

=== TARGET JOB ===
Title: {job.title}
Company: {job.company}
Location: {job.location}  |  Remote: {job.remote_allowed}
Industry: {job.industry}
Required Skills: {", ".join(job.requirements)}
Nice-to-Have: {", ".join(job.nice_to_have)}
Description:
{job.description}

Generate the full resume now. After the resume, add a ```json block with:
{{"tailoring_notes": "<explanation of specific tailoring choices made>"}}
"""

    def _build_cover_letter_prompt(
        self, profile: UserProfile, job: JobListing
    ) -> str:
        return f"""Write a cover letter for the following candidate and job.

=== CANDIDATE ===
Name: {profile.name}
Location: {profile.location}
Key Skills: {", ".join(profile.skills[:10])}
Experience: {profile.years_of_experience} years as {", ".join(profile.desired_roles[:3])}
Notable:
- Education: {profile.education[0].get("degree", "") if profile.education else "Not specified"}
- Recent role: {profile.work_history[0].get("title", "") if len(profile.work_history) > 0 else "Not specified"}
  at {profile.work_history[0].get("company", "") if len(profile.work_history) > 0 else ""}

=== JOB ===
Title: {job.title}
Company: {job.company}
Location: {job.location}
Industry: {job.industry}
Must-have skills: {", ".join(job.requirements)}
Description: {job.description}

Write the cover letter now. After the letter, add a ```json block with:
{{"tailoring_notes": "<why specific paragraphs/phrases were chosen>"}}
"""
