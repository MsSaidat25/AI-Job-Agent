"""PDF and DOCX export for resume templates."""
from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO


@dataclass
class ParsedSection:
    """A parsed section of the resume."""

    level: int  # 1 = h1, 2 = h2, 3 = h3
    title: str
    items: list[str]


def parse_markdown(content: str) -> tuple[str, list[ParsedSection]]:
    """Parse resume markdown into a name/title and sections."""
    lines = content.strip().splitlines()
    name = ""
    sections: list[ParsedSection] = []
    current: ParsedSection | None = None

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("## "):
            name = stripped[2:].strip()
        elif stripped.startswith("## "):
            if current:
                sections.append(current)
            current = ParsedSection(level=2, title=stripped[3:].strip(), items=[])
        elif stripped.startswith("### "):
            if current:
                sections.append(current)
            current = ParsedSection(level=3, title=stripped[4:].strip(), items=[])
        elif stripped:
            if current:
                current.items.append(stripped)
            elif not name:
                name = stripped
        elif current and current.items:
            current.items.append("")  # preserve paragraph breaks

    if current:
        sections.append(current)

    return name, sections


def _sanitise_latin1(text: str) -> str:
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _strip_bold(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def export_pdf(content_md: str, template_id: str = "classic") -> bytes:
    """Render resume markdown as a styled PDF using the specified template."""
    from fpdf import FPDF

    from src.resume_templates import get_template

    style = get_template(template_id)
    name, sections = parse_markdown(content_md)

    pdf = FPDF()
    pdf.set_margins(style.margin_left, style.margin_top, style.margin_right)
    pdf.set_auto_page_break(auto=True, margin=style.margin_bottom)
    pdf.add_page()

    page_w = pdf.w - style.margin_left - style.margin_right

    # Header
    if style.header_style == "banner":
        pdf.set_fill_color(*style.accent_color)
        pdf.rect(0, 0, pdf.w, 28, "F")
        pdf.set_text_color(255, 255, 255)
        pdf.set_font(style.heading_font, "B", style.font_size_name)
        pdf.set_y(6)
        pdf.cell(0, 12, _sanitise_latin1(name), align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.set_y(30)
        pdf.set_text_color(*style.primary_color)
    elif style.header_style == "centered":
        pdf.set_text_color(*style.accent_color)
        pdf.set_font(style.heading_font, "B", style.font_size_name)
        pdf.cell(0, 10, _sanitise_latin1(name), align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(*style.primary_color)
        pdf.ln(2)
    else:
        pdf.set_text_color(*style.accent_color)
        pdf.set_font(style.heading_font, "B", style.font_size_name)
        pdf.cell(0, 10, _sanitise_latin1(name), new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(*style.primary_color)
        pdf.ln(1)

    # Contact line extraction
    if sections and sections[0].items:
        contact_candidates = []
        remaining_items = []
        for item in sections[0].items:
            clean = _strip_bold(item)
            if "|" in clean or "@" in clean:
                contact_candidates.append(clean)
            else:
                remaining_items.append(item)
        if contact_candidates:
            pdf.set_font(style.body_font, "", style.font_size_body)
            pdf.set_text_color(100, 100, 100)
            align = "C" if style.header_style == "centered" else "L"
            for contact in contact_candidates:
                pdf.cell(
                    0, style.line_spacing, _sanitise_latin1(contact),
                    align=align, new_x="LMARGIN", new_y="NEXT",
                )
            pdf.set_text_color(*style.primary_color)
            pdf.ln(2)
            sections[0].items = remaining_items

    # Sections
    for section in sections:
        if style.section_divider == "line":
            y = pdf.get_y()
            pdf.set_draw_color(*style.accent_color)
            pdf.set_line_width(0.3)
            pdf.line(style.margin_left, y, pdf.w - style.margin_right, y)
            pdf.ln(2)
        elif style.section_divider == "dots":
            pdf.set_font(style.body_font, "", 6)
            pdf.set_text_color(*style.accent_color)
            pdf.cell(0, 3, (" . " * 40)[:80], new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(*style.primary_color)
            pdf.ln(1)
        elif style.section_divider == "space":
            pdf.ln(4)

        font_size = style.font_size_heading if section.level == 2 else style.font_size_subheading
        pdf.set_font(style.heading_font, "B", font_size)
        pdf.set_text_color(*style.accent_color)
        pdf.cell(0, 8, _sanitise_latin1(section.title.upper()), new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(*style.primary_color)
        pdf.ln(1)

        pdf.set_font(style.body_font, "", style.font_size_body)
        for item in section.items:
            if not item:
                pdf.ln(2)
                continue
            if item.startswith(("- ", "* ")):
                text = _strip_bold(item[2:])
                bullet = "\u2022 " if template_id != "ats_optimized" else "- "
                pdf.multi_cell(
                    page_w, style.line_spacing,
                    _sanitise_latin1(f"  {bullet}{text}"),
                    new_x="LMARGIN", new_y="NEXT",
                )
            else:
                text = _strip_bold(item)
                if "**" in item or ("|" in item and len(item) < 120):
                    pdf.set_font(style.body_font, "B", style.font_size_body)
                    pdf.multi_cell(page_w, style.line_spacing, _sanitise_latin1(text), new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font(style.body_font, "", style.font_size_body)
                else:
                    pdf.multi_cell(page_w, style.line_spacing, _sanitise_latin1(text), new_x="LMARGIN", new_y="NEXT")

    return bytes(pdf.output())


def export_docx(content_md: str, template_id: str = "classic") -> bytes:
    """Render resume markdown as a styled DOCX using the specified template."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    from src.resume_templates import get_template

    style = get_template(template_id)
    name, sections = parse_markdown(content_md)

    doc = Document()

    for sec in doc.sections:
        sec.top_margin = Inches(style.margin_top / 25.4)
        sec.bottom_margin = Inches(style.margin_bottom / 25.4)
        sec.left_margin = Inches(style.margin_left / 25.4)
        sec.right_margin = Inches(style.margin_right / 25.4)

    default_style = doc.styles["Normal"]
    default_style.font.size = Pt(style.font_size_body)  # type: ignore[union-attr]
    default_style.font.color.rgb = RGBColor(*style.primary_color)  # type: ignore[union-attr]
    pf = default_style.paragraph_format  # type: ignore[union-attr]
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(style.line_spacing + 8)

    name_para = doc.add_paragraph()
    if style.header_style == "centered":
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(name)
    run.bold = True
    run.font.size = Pt(style.font_size_name)
    run.font.color.rgb = RGBColor(*style.accent_color)

    # Contact info
    if sections and sections[0].items:
        contact_items = []
        remaining = []
        for item in sections[0].items:
            clean = _strip_bold(item)
            if "|" in clean or "@" in clean:
                contact_items.append(clean)
            else:
                remaining.append(item)
        if contact_items:
            for contact in contact_items:
                cp = doc.add_paragraph()
                if style.header_style == "centered":
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(contact)
                cr.font.size = Pt(style.font_size_body)
                cr.font.color.rgb = RGBColor(100, 100, 100)
            sections[0].items = remaining

    for section in sections:
        if style.section_divider == "line":
            border_para = doc.add_paragraph()
            border_para.paragraph_format.space_before = Pt(6)
            border_para.paragraph_format.space_after = Pt(2)
            pPr = border_para._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(
                qn("w:bottom"),
                {
                    qn("w:val"): "single",
                    qn("w:sz"): "4",
                    qn("w:space"): "1",
                    qn("w:color"): "{:02X}{:02X}{:02X}".format(*style.accent_color),
                },
            )
            pBdr.append(bottom)
            pPr.append(pBdr)

        heading_level = 1 if section.level == 2 else 2
        heading = doc.add_heading(section.title.upper(), level=heading_level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*style.accent_color)
            run.font.size = Pt(
                style.font_size_heading if section.level == 2 else style.font_size_subheading
            )

        for item in section.items:
            if not item:
                continue
            if item.startswith(("- ", "* ")):
                text = _strip_bold(item[2:])
                bp = doc.add_paragraph(style="List Bullet")
                br = bp.add_run(text)
                br.font.size = Pt(style.font_size_body)
            else:
                para = doc.add_paragraph()
                parts = re.split(r"\*\*(.+?)\*\*", item)
                for i, part in enumerate(parts):
                    run = para.add_run(part)
                    run.bold = i % 2 == 1
                    run.font.size = Pt(style.font_size_body)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
